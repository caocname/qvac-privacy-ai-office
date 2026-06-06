"""
知识库 API — 文档上传、解析、分片、向量化、版本管理、文件夹、下载。
"""
from __future__ import annotations

import os
import uuid
from datetime import datetime, timezone
from pathlib import Path

from fastapi import APIRouter, UploadFile, File, Form, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from backend.config import UPLOAD_DIR, RAG_CHUNK_SIZE, RAG_CHUNK_OVERLAP
from backend.database.connection import DatabaseManager
from backend.logger.audit_logger import get_audit_logger
from backend.logger.log_models import LogType
from backend.services.llm_service import get_llm_service
from backend.services.rag_service import RAGService
from backend.state_machine import StateCode, get_state_manager

router = APIRouter(prefix="/api/v1/knowledge", tags=["knowledge"])


def _extract_text_from_txt(file_path: Path) -> str:
    for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            return file_path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return ""


def _extract_text_from_pdf(file_path: Path) -> str:
    try:
        import fitz
        doc = fitz.open(str(file_path))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except ImportError:
        pass
    try:
        from pdfminer.high_level import extract_text
        return extract_text(str(file_path))
    except ImportError:
        pass
    return ""


def _extract_text_from_docx(file_path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        pass
    return ""


def _extract_text_from_md(file_path: Path) -> str:
    for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            return file_path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return ""


def _chunk_text(text: str, chunk_size: int = RAG_CHUNK_SIZE, overlap: int = RAG_CHUNK_OVERLAP) -> list[str]:
    paragraphs = text.split("\n")
    chunks = []
    current = ""
    current_tokens = 0

    for para in paragraphs:
        para_tokens = RAGService._estimate_tokens(para)
        if current_tokens + para_tokens > chunk_size and current:
            chunks.append(current.strip())
            current = para
            current_tokens = para_tokens
        else:
            current = (current + "\n" + para).strip() if current else para
            current_tokens += para_tokens

    if current.strip():
        chunks.append(current.strip())

    if len(chunks) <= 1 and len(text) > 500:
        chunks = []
        for i in range(0, len(text), chunk_size * 2):
            chunks.append(text[i:i + chunk_size * 3])

    return chunks


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.%f")[:-3] + "Z"


def _build_folder_tree(folders: list[dict], parent_id: str | None = None) -> list[dict]:
    """递归构建文件夹树。"""
    tree = []
    for f in folders:
        if f.get("parent_id") == parent_id:
            children = _build_folder_tree(folders, f["folder_id"])
            node = dict(f)
            if children:
                node["children"] = children
            tree.append(node)
    return tree


# ---- 文件夹 API ----

class CreateFolderRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=256)
    parent_id: str | None = None


@router.post("/folders")
async def create_folder(req: CreateFolderRequest):
    """创建文件夹。"""
    db = DatabaseManager.get_instance()
    folder_id = f"folder-{uuid.uuid4().hex[:12]}"
    now = _now_iso()
    db.conn.execute(
        "INSERT INTO knowledge_folders (folder_id, name, parent_id, create_time, update_time) VALUES (?, ?, ?, ?, ?)",
        (folder_id, req.name, req.parent_id, now, now),
    )
    db.conn.commit()
    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {"folder_id": folder_id, "name": req.name, "parent_id": req.parent_id, "create_time": now},
    }


@router.get("/folders")
async def list_folders(tree: bool = Query(default=True)):
    """列出文件夹（默认树形结构）。"""
    db = DatabaseManager.get_instance()
    rows = db.conn.execute(
        "SELECT folder_id, name, parent_id, create_time, update_time FROM knowledge_folders ORDER BY create_time ASC"
    ).fetchall()
    folders = [
        {"folder_id": r[0], "name": r[1], "parent_id": r[2], "create_time": r[3], "update_time": r[4]}
        for r in rows
    ]
    if tree:
        folders = _build_folder_tree(folders)
    return {"code": StateCode.IDLE.value, "status": StateCode.IDLE.name, "data": folders}


@router.put("/folders/{folder_id}")
async def rename_folder(folder_id: str, name: str = Query(..., min_length=1, max_length=256)):
    """重命名文件夹。"""
    db = DatabaseManager.get_instance()
    now = _now_iso()
    db.conn.execute(
        "UPDATE knowledge_folders SET name = ?, update_time = ? WHERE folder_id = ?",
        (name, now, folder_id),
    )
    db.conn.commit()
    return {"code": StateCode.IDLE.value, "status": StateCode.IDLE.name, "data": {"folder_id": folder_id, "name": name}}


@router.delete("/folders/{folder_id}")
async def delete_folder(folder_id: str):
    """删除文件夹（文件夹内文档的 folder_id 置空）。"""
    db = DatabaseManager.get_instance()
    db.conn.execute("UPDATE knowledge_base SET folder_id = NULL WHERE folder_id = ?", (folder_id,))
    db.conn.execute("DELETE FROM knowledge_folders WHERE folder_id = ?", (folder_id,))
    db.conn.commit()
    return {"code": StateCode.IDLE.value, "status": StateCode.IDLE.name, "message": "文件夹已删除"}


# ---- 文档上传（含版本管理） ----

@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    isolate_mode: str = Form(default="session"),
    session_id: str | None = Form(default=None),
    folder_id: str | None = Form(default=None),
):
    """知识库文件上传 — 支持重复导入版本管理。

    同一文件重复导入时自动创建新版本（version 递增），通过 import_group_id 关联。
    """
    state_mgr = get_state_manager()
    logger = get_audit_logger()
    file_path = None

    try:
        state_mgr.transition(StateCode.FILE_UPLOADING)

        if isolate_mode == "session" and not session_id:
            return {
                "code": StateCode.ERROR.value,
                "status": StateCode.ERROR.name,
                "error_class": "MISSING_SESSION_ID",
                "message": "session 隔离模式必须提供 session_id",
            }

        file_type = file.filename.rsplit(".", 1)[-1].lower() if file.filename else "txt"
        if file_type not in ("pdf", "docx", "txt", "md"):
            return {
                "code": StateCode.ERROR.value,
                "status": StateCode.ERROR.name,
                "error_class": "UNSUPPORTED_FILE_TYPE",
                "message": f"不支持的文件类型: {file_type}，仅支持 PDF、DOCX、TXT、MD",
            }

        db = DatabaseManager.get_instance()
        original_name = file.filename.strip()

        # 版本管理：检查是否已有同名文件（大小写不敏感 + 去空格）
        existing = db.conn.execute(
            "SELECT file_id, version, import_group_id FROM knowledge_base "
            "WHERE LOWER(original_name) = LOWER(?) AND is_deleted = 0 ORDER BY version DESC LIMIT 1",
            (original_name,),
        ).fetchone()

        if existing:
            import_group_id = existing[2] or f"group-{uuid.uuid4().hex[:12]}"
            new_version = existing[1] + 1
        else:
            import_group_id = f"group-{uuid.uuid4().hex[:12]}"
            new_version = 1

        file_id = f"file-{uuid.uuid4().hex[:16]}"
        UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        file_path = UPLOAD_DIR / f"{file_id}.{file_type}"

        content = await file.read()
        file_path.write_bytes(content)
        file_size = len(content)

        logger.log(LogType.KNOWLEDGE_UPLOAD, {
            "file_id": file_id,
            "file_name": file.filename,
            "file_type": file_type,
            "file_size": file_size,
            "isolate_mode": isolate_mode,
            "session_id": session_id,
            "folder_id": folder_id,
            "version": new_version,
            "import_group_id": import_group_id,
        }, f"Upload: {file.filename} v{new_version} ({file_size} bytes)")

        state_mgr.transition(StateCode.FILE_PROCESSING)

        if file_type == "txt":
            extracted_text = _extract_text_from_txt(file_path)
        elif file_type == "pdf":
            extracted_text = _extract_text_from_pdf(file_path)
        elif file_type == "docx":
            extracted_text = _extract_text_from_docx(file_path)
        elif file_type == "md":
            extracted_text = _extract_text_from_md(file_path)
        else:
            extracted_text = ""

        if not extracted_text.strip():
            try:
                os.remove(file_path)
            except OSError:
                pass
            state_mgr.transition(StateCode.IDLE)
            return {
                "code": StateCode.ERROR.value,
                "status": StateCode.ERROR.name,
                "error_class": "EXTRACTION_FAILED",
                "message": "文件内容提取失败，文件可能已损坏或为空。",
            }

        chunks = _chunk_text(extracted_text)
        total_pages = max(1, len(chunks))

        cipher = db.cipher
        encrypted_path = cipher.encrypt(str(file_path)) if cipher else str(file_path)

        db.conn.execute(
            "INSERT INTO knowledge_base (file_id, file_name, file_path, file_size, total_pages, "
            "isolate_mode, session_id, folder_id, version, original_name, import_group_id) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (file_id, file.filename, encrypted_path, file_size, total_pages,
             isolate_mode, session_id, folder_id, new_version, original_name, import_group_id),
        )
        db.conn.commit()

        state_mgr.set_worker(StateCode.EMBEDDING, True)

        llm = get_llm_service()
        embeddings = await llm.embed(chunks)

        indexed_count = 0
        embedding_ok = False
        if embeddings and len(embeddings) == len(chunks):
            import numpy as np
            rag = RAGService()
            indexed_count = rag.index_document(
                file_id=file_id,
                file_name=file.filename,
                chunks=chunks,
                embeddings=[np.array(e, dtype=np.float32) for e in embeddings],
            )
            embedding_ok = True
        else:
            # Embedding 模型未就绪 — 仍保留文档和文本分片（翻译/上下文可正常使用），向量索引待后续重建
            logger.log(LogType.ERROR, {
                "event": "embedding_skipped",
                "file_id": file_id,
                "file_name": file.filename,
                "chunk_count": len(chunks),
                "embed_loaded": llm.is_embed_loaded,
                "note": "FAISS 索引已跳过，文档文本已保存，可在 Embedding 模型就绪后通过重新上传来补索引",
            })
            for i, chunk_text in enumerate(chunks):
                chunk_id = f"chunk-{uuid.uuid4().hex[:12]}"
                db.conn.execute(
                    "INSERT INTO rag_chunks (chunk_id, file_id, chunk_index, content, token_count, faiss_vector_id) "
                    "VALUES (?, ?, ?, ?, ?, ?)",
                    (chunk_id, file_id, i, chunk_text, RAGService._estimate_tokens(chunk_text), -1),
                )
            db.conn.commit()

        state_mgr.set_worker(StateCode.EMBEDDING, False)
        state_mgr.transition(StateCode.IDLE)

        return {
            "code": StateCode.IDLE.value,
            "status": StateCode.IDLE.name,
            "data": {
                "file_id": file_id,
                "file_name": file.filename,
                "total_pages": total_pages,
                "chunk_count": len(chunks),
                "indexed_count": indexed_count,
                "isolate_mode": isolate_mode,
                "folder_id": folder_id,
                "version": new_version,
                "import_group_id": import_group_id,
                "embedding_ok": embedding_ok,
            },
        }

    except Exception as exc:
        state_mgr.transition(StateCode.ERROR)
        state_mgr.transition(StateCode.IDLE)
        logger.log(LogType.ERROR, {
            "event": "upload_exception",
            "file_name": file.filename if file else "unknown",
            "error_class": type(exc).__name__,
        }, str(exc))
        if file_path is not None:
            try:
                os.remove(file_path)
            except OSError:
                pass
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "error_class": type(exc).__name__,
            "message": f"上传处理异常: {str(exc)[:300]}",
        }


# ---- 文件列表 ----

@router.get("/list")
async def list_files(
    session_id: str | None = None,
    folder_id: str | None = None,
):
    """列出知识库文件，支持按文件夹筛选。"""
    db = DatabaseManager.get_instance()
    query = (
        "SELECT file_id, file_name, file_size, total_pages, isolate_mode, session_id, "
        "folder_id, version, original_name, import_group_id, is_deleted, create_time "
        "FROM knowledge_base WHERE is_deleted = 0 "
    )
    params: list = []

    if folder_id:
        query += "AND folder_id = ? "
        params.append(folder_id)

    if session_id:
        query += "AND (isolate_mode = 'global' OR session_id = ?) "
        params.append(session_id)

    query += "ORDER BY create_time DESC"
    rows = db.conn.execute(query, params).fetchall()

    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": [
            {
                "file_id": r[0], "file_name": r[1], "file_size": r[2],
                "total_pages": r[3], "isolate_mode": r[4], "session_id": r[5],
                "folder_id": r[6], "version": r[7], "original_name": r[8],
                "import_group_id": r[9], "is_deleted": r[10], "create_time": r[11],
            }
            for r in rows
        ],
    }


# ---- 版本管理 ----

@router.get("/versions/{import_group_id}")
async def list_versions(import_group_id: str):
    """列出同一文档的所有版本。"""
    db = DatabaseManager.get_instance()
    rows = db.conn.execute(
        "SELECT file_id, file_name, file_size, version, create_time "
        "FROM knowledge_base WHERE import_group_id = ? AND is_deleted = 0 "
        "ORDER BY version DESC",
        (import_group_id,),
    ).fetchall()
    if not rows:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "未找到该文档的版本记录",
        }
    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": [
            {
                "file_id": r[0], "file_name": r[1], "file_size": r[2],
                "version": r[3], "create_time": r[4],
            }
            for r in rows
        ],
    }


# ---- 文档下载（含时间戳） ----

@router.get("/download/{file_id}")
async def download_file(file_id: str):
    """下载指定版本的文档，文件名自动添加时间戳。

    导出文件名格式: {原名}_v{版本号}_{导入时间}.{扩展名}
    """
    db = DatabaseManager.get_instance()
    row = db.conn.execute(
        "SELECT file_name, file_path, version, create_time FROM knowledge_base "
        "WHERE file_id = ? AND is_deleted = 0",
        (file_id,),
    ).fetchone()
    if not row:
        from fastapi.responses import JSONResponse
        return JSONResponse(
            status_code=404,
            content={"code": StateCode.ERROR.value, "status": StateCode.ERROR.name, "message": "文档不存在或已删除"},
        )

    file_name, encrypted_path, version, create_time = row
    cipher = db.cipher
    try:
        actual_path_str = cipher.decrypt(encrypted_path) if cipher else encrypted_path
    except Exception:
        actual_path_str = encrypted_path

    actual_path = Path(actual_path_str)
    if not actual_path.exists():
        from fastapi.responses import JSONResponse
        return JSONResponse(
            status_code=404,
            content={"code": StateCode.ERROR.value, "status": StateCode.ERROR.name, "message": "文件已从磁盘删除"},
        )

    # 构建带时间戳的下载文件名
    base_name = Path(file_name).stem
    ext = Path(file_name).suffix
    time_part = create_time.replace(":", "").replace("-", "").replace("T", "_").replace("Z", "")[:15] if create_time else ""
    download_name = f"{base_name}_v{version}_{time_part}{ext}"

    return FileResponse(
        path=str(actual_path),
        filename=download_name,
        media_type="application/octet-stream",
    )


# ---- 删除 ----

@router.delete("/delete")
async def delete_file(file_id: str):
    """三步物理删除：is_deleted=1 → 异步 os.remove → 冷启动清扫"""
    db = DatabaseManager.get_instance()
    row = db.conn.execute(
        "SELECT file_path FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
        (file_id,),
    ).fetchone()
    if not row:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文件不存在或已删除",
        }

    db.conn.execute("UPDATE knowledge_base SET is_deleted = 1 WHERE file_id = ?", (file_id,))
    db.conn.execute("DELETE FROM rag_chunks WHERE file_id = ?", (file_id,))
    db.conn.commit()

    cipher = db.cipher
    encrypted_path = row[0]
    try:
        actual_path = cipher.decrypt(encrypted_path) if cipher else encrypted_path
    except Exception:
        actual_path = encrypted_path

    def _erase_file(path):
        try:
            os.remove(path)
        except OSError:
            pass

    from concurrent.futures import ThreadPoolExecutor
    ThreadPoolExecutor(max_workers=1).submit(_erase_file, actual_path)

    get_audit_logger().log(LogType.KNOWLEDGE_DELETE, {
        "file_id": file_id,
        "event": "logical_delete_committed",
    })

    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "message": "文件已标记删除，后台异步擦除中。",
    }


# ---- 移动文件到文件夹 ----

class MoveFileRequest(BaseModel):
    file_id: str = Field(..., description="文件 ID")
    folder_id: str | None = Field(default=None, description="目标文件夹 ID，为空则移出文件夹")


@router.put("/files/move")
async def move_file(req: MoveFileRequest):
    """将文件移动到指定文件夹（folder_id 为空则移出文件夹）。"""
    db = DatabaseManager.get_instance()
    row = db.conn.execute(
        "SELECT file_id FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
        (req.file_id,),
    ).fetchone()
    if not row:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文件不存在或已删除",
        }

    # 验证目标文件夹存在
    if req.folder_id:
        folder = db.conn.execute(
            "SELECT folder_id FROM knowledge_folders WHERE folder_id = ?",
            (req.folder_id,),
        ).fetchone()
        if not folder:
            return {
                "code": StateCode.ERROR.value,
                "status": StateCode.ERROR.name,
                "message": "目标文件夹不存在",
            }

    db.conn.execute(
        "UPDATE knowledge_base SET folder_id = ? WHERE file_id = ?",
        (req.folder_id, req.file_id),
    )
    db.conn.commit()

    get_audit_logger().log(LogType.SYSTEM, {
        "event": "file_moved",
        "file_id": req.file_id,
        "folder_id": req.folder_id,
    })

    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {"file_id": req.file_id, "folder_id": req.folder_id},
    }


# ---- 文档内容 ----

@router.get("/document/{file_id}/content")
async def get_document_content(file_id: str):
    """获取文档全文 — 从 rag_chunks 按序拼接。"""
    db = DatabaseManager.get_instance()

    row = db.conn.execute(
        "SELECT file_name FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
        (file_id,),
    ).fetchone()
    if not row:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档不存在或已删除",
        }

    chunks = db.conn.execute(
        "SELECT content FROM rag_chunks WHERE file_id = ? ORDER BY chunk_index",
        (file_id,),
    ).fetchall()

    if not chunks:
        path_row = db.conn.execute(
            "SELECT file_path FROM knowledge_base WHERE file_id = ?", (file_id,)
        ).fetchone()
        if path_row:
            cipher = db.cipher
            enc_path = path_row[0]
            try:
                actual_path = cipher.decrypt(enc_path) if cipher else enc_path
            except Exception:
                actual_path = enc_path
            file_path = Path(actual_path)
            if file_path.exists():
                ext = file_path.suffix.lower()
                if ext == ".txt":
                    extracted = _extract_text_from_txt(file_path)
                elif ext == ".pdf":
                    extracted = _extract_text_from_pdf(file_path)
                elif ext == ".docx":
                    extracted = _extract_text_from_docx(file_path)
                elif ext == ".md":
                    extracted = _extract_text_from_md(file_path)
                else:
                    extracted = ""
                if extracted.strip():
                    return {
                        "code": StateCode.IDLE.value,
                        "status": StateCode.IDLE.name,
                        "data": {
                            "file_id": file_id,
                            "file_name": row[0],
                            "content": extracted,
                            "char_count": len(extracted),
                        },
                    }
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档内容为空，请重新上传。",
        }

    full_text = "\n\n".join(c[0] for c in chunks)
    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {
            "file_id": file_id,
            "file_name": row[0],
            "content": full_text,
            "char_count": len(full_text),
        },
    }


# ---- 导出格式转换 ----

class ExportRequest(BaseModel):
    file_id: str = Field(..., description="文档 ID")
    format: str = Field(default="txt", description="导出格式: txt | md | docx")


def _convert_to_md(text: str, title: str = "") -> str:
    """纯文本 → Markdown 格式转换。"""
    lines = text.split("\n")
    md_lines = []
    if title:
        md_lines.append(f"# {title}\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            md_lines.append("")
        elif stripped.startswith("#"):
            md_lines.append(stripped)
        elif len(stripped) < 80 and stripped.endswith(("：", ":")) and not stripped.startswith("-"):
            md_lines.append(f"## {stripped}")
        else:
            md_lines.append(stripped)
    return "\n\n".join(md_lines)


def _convert_to_docx(text: str, title: str = "") -> bytes:
    """纯文本 → DOCX 格式转换。"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        doc = DocxDocument()
        if title:
            h = doc.add_heading(title, level=1)
            h.runs[0].font.size = Pt(16) if h.runs else None
        for para_text in text.split("\n"):
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)
        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except ImportError:
        return b""


@router.post("/export")
async def export_document(req: ExportRequest):
    """导出文档为指定格式（txt / md / docx）。"""
    db = DatabaseManager.get_instance()
    row = db.conn.execute(
        "SELECT file_name FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
        (req.file_id,),
    ).fetchone()
    if not row:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档不存在或已删除",
        }

    file_name = row[0]
    base_name = Path(file_name).stem

    # 获取全文
    chunks = db.conn.execute(
        "SELECT content FROM rag_chunks WHERE file_id = ? ORDER BY chunk_index",
        (req.file_id,),
    ).fetchall()
    full_text = "\n\n".join(c[0] for c in chunks) if chunks else ""

    fmt = req.format.lower()
    if fmt == "md":
        output = _convert_to_md(full_text, base_name)
        download_name = f"{base_name}.md"
        media_type = "text/markdown"
        import tempfile
        tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8")
        tmp.write(output)
        tmp.close()
        return FileResponse(path=tmp.name, filename=download_name, media_type=media_type)

    elif fmt == "docx":
        output_bytes = _convert_to_docx(full_text, base_name)
        download_name = f"{base_name}.docx"
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.write(output_bytes)
        tmp.close()
        return FileResponse(path=tmp.name, filename=download_name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    else:  # txt
        download_name = f"{base_name}.txt"
        import tempfile
        tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8")
        tmp.write(full_text)
        tmp.close()
        return FileResponse(path=tmp.name, filename=download_name, media_type="text/plain")


# ---- 单文件隔离模式切换 ----

class ChangeIsolateRequest(BaseModel):
    file_id: str = Field(..., description="文档 ID")
    isolate_mode: str = Field(..., description="新隔离模式: global | session | temp")
    session_id: str | None = Field(default=None)


@router.put("/files/isolate")
async def change_isolate_mode(req: ChangeIsolateRequest):
    """修改单个文档的隔离模式。"""
    db = DatabaseManager.get_instance()
    row = db.conn.execute(
        "SELECT file_id FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
        (req.file_id,),
    ).fetchone()
    if not row:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档不存在或已删除",
        }
    if req.isolate_mode not in ("global", "session", "temp"):
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "无效的隔离模式，仅支持 global / session / temp",
        }

    db.conn.execute(
        "UPDATE knowledge_base SET isolate_mode = ?, session_id = ? WHERE file_id = ?",
        (req.isolate_mode, req.session_id, req.file_id),
    )
    db.conn.commit()

    get_audit_logger().log(LogType.SYSTEM, {
        "event": "isolate_changed",
        "file_id": req.file_id,
        "new_mode": req.isolate_mode,
    })
    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {"file_id": req.file_id, "isolate_mode": req.isolate_mode},
    }


# ---- 批量操作 ----

class BatchDeleteRequest(BaseModel):
    file_ids: list[str] = Field(..., description="要删除的文档 ID 列表")


@router.post("/batch/delete")
async def batch_delete(req: BatchDeleteRequest):
    """批量删除文档。"""
    db = DatabaseManager.get_instance()
    deleted = []
    for file_id in req.file_ids:
        row = db.conn.execute(
            "SELECT file_path FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
            (file_id,),
        ).fetchone()
        if not row:
            continue
        db.conn.execute("UPDATE knowledge_base SET is_deleted = 1 WHERE file_id = ?", (file_id,))
        db.conn.execute("DELETE FROM rag_chunks WHERE file_id = ?", (file_id,))
        cipher = db.cipher
        enc_path = row[0]
        try:
            actual = cipher.decrypt(enc_path) if cipher else enc_path
        except Exception:
            actual = enc_path

        def _erase(p):
            try:
                os.remove(p)
            except OSError:
                pass

        from concurrent.futures import ThreadPoolExecutor
        ThreadPoolExecutor(max_workers=1).submit(_erase, actual)
        deleted.append(file_id)

    db.conn.commit()

    get_audit_logger().log(LogType.KNOWLEDGE_DELETE, {
        "event": "batch_delete",
        "count": len(deleted),
        "file_ids": deleted,
    })
    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {"deleted_count": len(deleted), "deleted_ids": deleted},
    }


class BatchIsolateRequest(BaseModel):
    file_ids: list[str] = Field(..., description="文档 ID 列表")
    isolate_mode: str = Field(..., description="隔离模式: global | session | temp")
    session_id: str | None = Field(default=None)


@router.post("/batch/isolate")
async def batch_isolate(req: BatchIsolateRequest):
    """批量修改文档隔离模式。"""
    if req.isolate_mode not in ("global", "session", "temp"):
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "无效的隔离模式",
        }
    db = DatabaseManager.get_instance()
    updated = []
    for file_id in req.file_ids:
        row = db.conn.execute(
            "SELECT file_id FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
            (file_id,),
        ).fetchone()
        if not row:
            continue
        db.conn.execute(
            "UPDATE knowledge_base SET isolate_mode = ?, session_id = ? WHERE file_id = ?",
            (req.isolate_mode, req.session_id, file_id),
        )
        updated.append(file_id)

    db.conn.commit()
    get_audit_logger().log(LogType.SYSTEM, {
        "event": "batch_isolate",
        "count": len(updated),
        "new_mode": req.isolate_mode,
    })
    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {"updated_count": len(updated), "isolate_mode": req.isolate_mode},
    }


# ---- 翻译 ----

class TranslateRequest(BaseModel):
    file_id: str = Field(..., description="文档 ID")
    source_lang: str = Field(default="auto", description="源语言")
    target_lang: str = Field(default="zh", description="目标语言")


@router.post("/translate")
async def translate_document(req: TranslateRequest):
    """翻译文档全文，通过 Bridge LLM 执行。"""
    import httpx
    from backend.services.llm_service import get_llm_service

    db = DatabaseManager.get_instance()
    row = db.conn.execute(
        "SELECT file_name FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
        (req.file_id,),
    ).fetchone()
    if not row:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档不存在或已删除",
        }

    # 获取全文
    chunks = db.conn.execute(
        "SELECT content FROM rag_chunks WHERE file_id = ? ORDER BY chunk_index",
        (req.file_id,),
    ).fetchall()
    if not chunks:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档内容为空",
        }

    # 检查 Bridge LLM 是否就绪
    llm_svc = get_llm_service()
    bridge_ok = await llm_svc.ping()
    if not bridge_ok:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "AI 服务尚未就绪，请等待模型加载完成后重试（启动后约需 30-60 秒）",
        }
    await llm_svc.health()
    if not llm_svc.is_llm_loaded:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "LLM 模型尚未加载完成，请稍后重试",
        }

    full_text = "\n\n".join(c[0] for c in chunks)
    file_name = row[0]

    # 按 ~1000 字符分段翻译（避免 1B 小模型上下文溢出导致输出崩溃）
    part_size = 1000
    parts = [full_text[i:i + part_size] for i in range(0, len(full_text), part_size)]

    translated_parts = []
    async with httpx.AsyncClient(timeout=300.0) as client:
        for i, part in enumerate(parts):
            try:
                resp = await client.post(
                    "http://127.0.0.1:18889/api/translate",
                    json={
                        "text": part,
                        "source_lang": req.source_lang,
                        "target_lang": req.target_lang,
                    },
                )
                if resp.status_code == 200:
                    translated_parts.append(resp.json().get("translated_text", ""))
                else:
                    translated_parts.append(f"[翻译失败] {part[:100]}...")
            except Exception as exc:
                translated_parts.append(f"[翻译异常: {str(exc)[:100]}]")

    translated_text = "\n\n".join(translated_parts)

    get_audit_logger().log(LogType.SYSTEM, {
        "event": "document_translated",
        "file_id": req.file_id,
        "file_name": file_name,
        "source_lang": req.source_lang,
        "target_lang": req.target_lang,
        "char_count": len(translated_text),
    })

    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {
            "file_id": req.file_id,
            "file_name": file_name,
            "original_text": full_text,
            "translated_text": translated_text,
            "source_lang": req.source_lang,
            "target_lang": req.target_lang,
        },
    }


class BatchTranslateRequest(BaseModel):
    file_ids: list[str] = Field(..., description="文档 ID 列表")
    source_lang: str = Field(default="auto")
    target_lang: str = Field(default="zh")


@router.post("/batch/translate")
async def batch_translate(req: BatchTranslateRequest):
    """批量翻译文档 — 异步并发执行。"""
    import asyncio

    async def _translate_one(file_id: str) -> dict:
        try:
            tr = TranslateRequest(
                file_id=file_id,
                source_lang=req.source_lang,
                target_lang=req.target_lang,
            )
            result = await translate_document(tr)
            if result.get("code") == 100:
                return {"file_id": file_id, "status": "ok", "translated_text": result["data"]["translated_text"]}
            return {"file_id": file_id, "status": "error", "message": result.get("message", "")}
        except Exception as exc:
            return {"file_id": file_id, "status": "error", "message": str(exc)[:200]}

    tasks = [_translate_one(fid) for fid in req.file_ids]
    results = await asyncio.gather(*tasks)

    return {
        "code": StateCode.IDLE.value,
        "status": StateCode.IDLE.name,
        "data": {"results": results},
    }


# ---- TTS 代理 ----

class TTSRequest(BaseModel):
    file_id: str = Field(..., description="文档 ID")
    language: str = Field(default="zh")


@router.post("/tts")
async def tts_document(req: TTSRequest):
    """获取文档内容的 TTS 音频 — 代理到 Bridge。

    返回 WAV 音频流或错误提示。
    """
    import httpx
    from fastapi.responses import StreamingResponse

    db = DatabaseManager.get_instance()
    row = db.conn.execute(
        "SELECT file_name FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
        (req.file_id,),
    ).fetchone()
    if not row:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档不存在或已删除",
        }

    # 获取全文
    chunks = db.conn.execute(
        "SELECT content FROM rag_chunks WHERE file_id = ? ORDER BY chunk_index",
        (req.file_id,),
    ).fetchall()
    full_text = "\n\n".join(c[0] for c in chunks) if chunks else ""
    if not full_text.strip():
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "文档内容为空",
        }

    # 限制长度（TTS 处理长文本很慢，取前 5000 字符）
    tts_text = full_text[:5000]

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(
                "http://127.0.0.1:18889/api/tts/speak",
                json={"text": tts_text, "language": req.language},
            )
            if resp.status_code == 200:
                audio_data = resp.content
                return StreamingResponse(
                    iter([audio_data]),
                    media_type="audio/wav",
                    headers={"Content-Disposition": f"attachment; filename={row[0]}.wav"},
                )
            else:
                detail = resp.json() if resp.headers.get("content-type") == "application/json" else {}
                return {
                    "code": StateCode.ERROR.value,
                    "status": StateCode.ERROR.name,
                    "message": f"TTS 失败: {detail.get('error', resp.text)[:200]}",
                }
    except Exception as exc:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": f"TTS 服务不可用: {str(exc)[:200]}",
        }


# ---- 批量导出 ----

class BatchExportRequest(BaseModel):
    file_ids: list[str] = Field(..., description="文档 ID 列表")
    format: str = Field(default="txt", description="导出格式: txt | md | docx")


@router.post("/batch/export")
async def batch_export(req: BatchExportRequest):
    """批量导出文档 — 打包为 ZIP。"""
    import tempfile
    import zipfile
    from fastapi.responses import FileResponse as FR

    db = DatabaseManager.get_instance()
    tmp_dir = tempfile.mkdtemp()
    exported = 0

    for file_id in req.file_ids:
        row = db.conn.execute(
            "SELECT file_name FROM knowledge_base WHERE file_id = ? AND is_deleted = 0",
            (file_id,),
        ).fetchone()
        if not row:
            continue

        file_name = row[0]
        base_name = Path(file_name).stem

        chunks = db.conn.execute(
            "SELECT content FROM rag_chunks WHERE file_id = ? ORDER BY chunk_index",
            (file_id,),
        ).fetchall()
        full_text = "\n\n".join(c[0] for c in chunks) if chunks else ""
        if not full_text:
            continue

        fmt = req.format.lower()
        if fmt == "md":
            content = _convert_to_md(full_text, base_name)
            ext = "md"
            fpath = os.path.join(tmp_dir, f"{base_name}.{ext}")
            with open(fpath, "w", encoding="utf-8") as f:
                f.write(content)
        elif fmt == "docx":
            content_bytes = _convert_to_docx(full_text, base_name)
            fpath = os.path.join(tmp_dir, f"{base_name}.docx")
            with open(fpath, "wb") as f:
                f.write(content_bytes)
        else:
            content = full_text
            ext = "txt"
            fpath = os.path.join(tmp_dir, f"{base_name}.{ext}")
            with open(fpath, "w", encoding="utf-8") as f:
                f.write(content)
        exported += 1

    if exported == 0:
        return {
            "code": StateCode.ERROR.value,
            "status": StateCode.ERROR.name,
            "message": "没有找到有效的文档内容",
        }

    zip_path = os.path.join(tmp_dir, "knowledge_export.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname in os.listdir(tmp_dir):
            fpath = os.path.join(tmp_dir, fname)
            if fname != "knowledge_export.zip":
                zf.write(fpath, fname)

    get_audit_logger().log(LogType.SYSTEM, {
        "event": "knowledge_batch_export",
        "count": exported,
        "format": fmt,
    })

    return FR(path=zip_path, filename="knowledge_batch_export.zip", media_type="application/zip")
